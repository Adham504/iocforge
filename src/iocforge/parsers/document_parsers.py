"""Parsers for binary document formats: PDF, DOCX, and ZIP archives."""
from __future__ import annotations

import tempfile
import zipfile
from pathlib import Path

from iocforge.parsers.base import BaseParser, ParsedDocument
from iocforge.utils.logger import get_logger

logger = get_logger(__name__)


class PDFParser(BaseParser):
    """Extract text from PDF files using pypdf."""

    extensions = (".pdf",)

    def parse(self, path: Path) -> ParsedDocument:
        text = ""
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
        except Exception as exc:  # pragma: no cover - depends on optional dep
            logger.error("Failed to parse PDF %s: %s", path, exc)
        return ParsedDocument(source=str(path), text=text)


class DOCXParser(BaseParser):
    """Extract text from Word .docx files using python-docx."""

    extensions = (".docx",)

    def parse(self, path: Path) -> ParsedDocument:
        text = ""
        try:
            import docx  # type: ignore

            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            text = "\n".join(parts)
        except Exception as exc:  # pragma: no cover - depends on optional dep
            logger.error("Failed to parse DOCX %s: %s", path, exc)
        return ParsedDocument(source=str(path), text=text)


class ZIPParser(BaseParser):
    """Recursively parse supported files inside a ZIP archive."""

    extensions = (".zip",)

    def __init__(self, factory: "object") -> None:
        # Lazily injected ParserFactory to avoid a circular import.
        self._factory = factory

    def parse(self, path: Path) -> ParsedDocument:
        root = ParsedDocument(source=str(path), text="")
        try:
            with zipfile.ZipFile(path) as archive:
                with tempfile.TemporaryDirectory() as tmp:
                    tmp_path = Path(tmp)
                    for name in archive.namelist():
                        if name.endswith("/"):
                            continue
                        try:
                            extracted = Path(archive.extract(name, tmp_path))
                        except Exception as exc:  # pragma: no cover
                            logger.warning("Skipping %s in zip: %s", name, exc)
                            continue
                        if self._factory.supports(extracted):  # type: ignore[attr-defined]
                            child = self._factory.parse(extracted)  # type: ignore[attr-defined]
                            # Re-label the source to reflect the archive path.
                            child.source = f"{path}::{name}"
                            root.children.append(child)
        except zipfile.BadZipFile as exc:
            logger.error("Bad ZIP file %s: %s", path, exc)
        return root
